"""
RAG 知识库服务 青春版 - Qwen3-Embedding-0.6B Embedding + ChromaDB 向量检索
"""
# 已支持从数据库与 backend/uploads 动态加载知识文档，并可在线增删改。
# TODO 后续可增加 RAG 生成式检索（RAG-Gen）
# TODO 后续可增加用户画像和会话上下文感知的个性化检索
# TODO 后续可增加多模态检索（图像内容 + 文本查询）
from __future__ import annotations

import os
import csv
import hashlib
import json
import re
import time
from pathlib import Path
from threading import RLock
import asyncio

from app.core.config import settings
from app.core.logging import brief_text, elapsed_ms, get_logger

logger = get_logger(__name__)

# 灵山胜境知识库(战损测试版)
LINGSHAN_KNOWLEDGE = [
    {
        "id": "ls_history_001",
        "category": "history",
        "title": "灵山胜境历史渊源",
        "content": """灵山胜境坐落于江苏省无锡市太湖西北部的马山镇，地处秦履峰、青龙山、白虎山三山环抱之间，
占地面积约30万平方米，是国家5A级旅游景区、世界佛教论坛永久会址，被誉为"东方佛国"和"太湖佛国"。
唐贞观年间，玄奘法师西行取经归来，见此地"层峦丛翠，曲水净秀，山形酷似印度灵鹫山"，
遂命名为"小灵山"，并嘱咐大弟子窥基法师在此住持道场，建小灵山庵。
北宋大中祥符年间（1008-1016年），宋真宗赐额"祥符禅寺"，成为江南名刹。
1994年修建工程奠基，1997年11月15日灵山大佛落成开光，2009年灵山梵宫正式开放。"""
    },
    {
        "id": "ls_dafa_001",
        "category": "attraction",
        "title": "灵山大佛",
        "content": """灵山大佛通高88米（佛体79米，莲花瓣9米），含台基总高101.5米，总用铜量725吨，
由2000块铸铜面板拼接而成，是世界最高露天青铜释迦牟尼立像。
右手施无畏印（除却众生痛苦），左手施与愿印（赐予众生欢乐）。
216级登云道暗合108烦恼与108愿望，前段108级"烦恼尽除"，后段108级"願望圆满"。
登顶可俯瞰太湖全景，夕阳时分金光普照，美不胜收。
开放时间：08:00-17:00（冬季提前至16:30）。"""
    },
    {
        "id": "ls_jiulong_001",
        "category": "performance",
        "title": "九龙灌浴表演",
        "content": """九龙灌浴总高27.2米，核心为7.2米高鎏金太子佛像，重12吨，周围环绕9条飞龙。
表演时莲花铜雕缓缓绽放，太子佛在《佛之诞》乐曲中升起旋转，九条飞龙同时喷出水柱为太子沐浴，
水幕与阳光交织出七彩佛光，完美再现"花开见佛"的祥瑞场景。
表演后可在广场两侧接取龙头流出的"圣水"，寓意祈福安康。
平日演出时间：10:00、11:30、13:30、15:00；
周末及节假日增加场次，每场约15分钟，建议提前10分钟到场占位。"""
    },
    {
        "id": "ls_fangong_001",
        "category": "attraction",
        "title": "灵山梵宫",
        "content": """灵山梵宫建筑面积72000㎡，最高处66.5米，被誉为"东方卢浮宫"，
第二、四届世界佛教论坛举办地，荣获中国建筑工程最高奖鲁班奖。
内部汇集东阳木雕、琉璃巨制《华藏世界》（160块彩色琉璃拼接，目前世界最大琉璃艺术作品之一）、
28米高星空穹顶、大型油画"世界佛教传法图"等非遗艺术。
圣坛为曼陀罗形态，可容纳2000人，全球唯一大型旋转舞台。
《灵山吉祥颂》演出：10:35、11:30、14:00、16:00，每场约20分钟，凭景区门票免费入场，建议提前30分钟排队。
开放时间：09:00-17:00（冬季16:30）。"""
    },
    {
        "id": "ls_wuyin_001",
        "category": "attraction",
        "title": "五印坛城",
        "content": """五印坛城位于香水海中央的独立圆岛上，被称为"小布达拉宫"。
五层重檐楼宇，总高约30米，藏式碉楼建筑风格，白墙红边金顶。
四门分别安置马宝、孔雀、共命鸟、象宝四尊瑞兽雕塑。
内部墙体绘有彩色唐卡（天然矿物颜料），转经筒长廊环绕主殿，108个纯铜转经筒。
顺时针转动转经筒，寓意祈福消灾、积累功德。
登至顶层观景台可俯瞰香水海、灵山梵宫与灵山大佛全景。
藏香制作体验：10:00、14:00（需预约，费用自理）。
开放时间：09:00-17:00（冬季16:30）。"""
    },
    {
        "id": "ls_xiangfu_001",
        "category": "temple",
        "title": "祥符禅寺",
        "content": """祥符禅寺始建于唐贞观年间，由玄奘法师弟子窥基大师开坛讲经，
北宋年间正式更名为"祥符禅寺"，是江南千年禅宗祖庭。
寺内有重12.8吨"祥符禅钟"（江南第一钟），钟声浑厚洪亮，响彻整个灵山山谷。
六角井是唐代名泉，曾被茶圣陆羽品鉴，列为江南名泉之一，井水清澈甘甜。
千年古银杏树龄超千年，秋季金黄的树叶铺满寺院，意境绝美。
可参与撞钟祈福，体验佛教文化的庄严与神圣。全天开放。"""
    },
    {
        "id": "ls_route_history",
        "category": "route",
        "title": "历史文化路线推荐",
        "content": """历史文化爱好者路线（约6小时深度游）：
南门入园 → 灵山大照壁（华夏第一壁，赵朴初题字）→ 佛手广场（天下第一掌）
→ 祥符禅寺（千年古刹，撞钟祈福）→ 灵山大佛（登216级台阶，俯瞰太湖）
→ 灵山梵宫（佛教艺术殿堂，观《吉祥颂》）→ 五印坛城（藏传佛教文化体验）→ 出口
讲解重点：玄奘法师与小灵山渊源、江南第一钟文化意义、青铜铸造工艺、
穹顶天象图创作依据、108转经筒祈福文化。
建议上午9点前入园，梵宫演出提前30分钟排队，夕阳时分拍大佛最美。"""
    },
    {
        "id": "ls_route_nature",
        "category": "route",
        "title": "自然风光路线推荐",
        "content": """自然风光爱好者路线（约5小时）：
南门入园 → 佛足坛（朝圣打卡）→ 九龙灌浴（观赏动态表演）
→ 菩提大道（250米印度菩提树拱廊）→ 灵山大佛（登顶俯瞰太湖）
→ 曼飞龙塔（傣族风格，夜间灯光绝美）→ 灵山精舍（禅意素斋）→ 出口
九龙灌浴提前10分钟到场，大佛平台观日落最佳，灵山精舍素斋值得品尝。"""
    },
    {
        "id": "ls_route_family",
        "category": "route",
        "title": "亲子家庭路线推荐",
        "content": """亲子家庭路线（约4小时轻松游）：
南门入园 → 九龙灌浴（动态表演，孩子最爱）→ 佛手广场（摸天下第一掌）
→ 百子戏弥勒（摸弥勒肚皮祈福，亲子互动拍照）→ 灵山梵宫（全息演出）
→ 五印坛城（转经筒体验）→ 出口
全程约4公里，孩子可乘观光车（40元/人），百子戏弥勒亲子互动超有趣，
梵宫《吉祥颂》演出孩子喜欢，建议提前30分钟占座。"""
    },
    {
        "id": "ls_ticket_001",
        "category": "faq",
        "title": "门票与优惠政策",
        "content": """灵山胜境门票价格：
成人票：210元（18周岁以上）
半价票：105元（6-18周岁未成年人、全日制本科及以下学生、60-69周岁老人）
免费：6周岁以下或1.4米以下儿童、70周岁以上老人、现役军人、残疾人
网购联票：225元（门票+观光车，无限次乘坐，更划算）
观光车单独购票：40元/人
导游服务：300元起"""
    },
    {
        "id": "ls_open_001",
        "category": "faq",
        "title": "开放时间与注意事项",
        "content": """开放时间：
夏季（4月-10月）：08:00-18:00
冬季（11月-3月）：08:30-17:00
建议上午9点前入园避开人流高峰，下午可观赏太湖日落。
最佳游览季节：春秋季节（3-5月、9-11月）气候宜人，春有樱花桃花，秋有银杏金黄。
注意事项：穿舒适运动鞋，夏季防晒，景区为佛教场所需保持安静，
不触摸佛像，部分区域禁止拍照（梵宫内禁用闪光灯）。"""
    },
    {
        "id": "ls_dining_001",
        "category": "dining",
        "title": "餐饮与住宿推荐",
        "content": """景区内餐饮：
梵宫素斋自助：50元/位，清淡雅致，体验佛门饮食文化，菜品丰富
素面套餐：35元/位，景区内多处餐厅，口味清淡，快速用餐
灵山精舍素斋：环境优雅，菜品精致，适合深度体验佛教文化
住宿：灵山精舍（景区内禅意酒店），含素斋与早课体验；
周边马山镇有多家酒店、民宿，价格几百至上千元不等。"""
    },
    {
        "id": "ls_nianhewan_001",
        "category": "nianhewan",
        "title": "拈花湾禅意小镇",
        "content": """拈花湾禅意小镇与灵山胜境比邻，以"禅意慢生活"为核心，开放时间09:00-21:30（冬季至20:30）。
主要景点：
拈花广场：入口核心，12米"拈花微笑"青铜雕塑，每日9:30开园仪式
梵天花海：30000㎡四季花海，春格桑花、夏硫华菊、秋波斯菊
香月花街：800米禅意商业街，非遗手作、禅茶美食、夜间灯笼绝美
拈花堂：免费禅坐冥想、抄经、禅茶品鉴，禅意讲座10:30/15:30
五灯湖：小镇最大水景，夜间《禅行》灯光秀19:00、20:00（约30分钟）"""
    },
    {
        "id": "ls_spots_other",
        "category": "attraction",
        "title": "其他特色景点",
        "content": """佛手广场（天下第一掌）：灵山大佛右手复制，高11.7米，宽5.5米，摸掌祈福保平安。
百子戏弥勒：9吨青铜群雕，弥勒佛身上百名孩童，摸肚皮寓意享一生福气。
曼飞龙塔：傣族风格九塔组合，主塔高16.9m，南传佛教代表，夜间灯光亮化绝美。
佛教文化博览馆：设于大佛三层座基内，10000㎡三层展馆，
一层五方五佛，二层世界佛教史，三层万佛殿（9999尊小佛像），免费参观。
无尽意斋：赵朴初先生纪念馆，北京四合院风格，展示书法作品与灵山渊源，免费开放。"""
    },
]


# Qwen Embedding 函数（供 ChromaDB 使用）
class QwenEmbeddingFunction:
    """
    修改后的 Embedding 函数：
    """
    def __init__(self, service: 'RAGService'):
        self.service = service

    def __call__(self, input: list[str]) -> list[list[float]]:
        # ChromaDB 的内置接口是同步的，这里使用内部 service 的同步包装方法
        return asyncio.run(self.service.embed_batch(input))

class MyLocalQwenEF:
    def __init__(self, model):
        self.model = model
    def __call__(self, input: list[str]) -> list[list[float]]:
        return self.model.encode(input).tolist()

class RAGService:
    """
    RAG 检索服务。

    知识来源不再是文件内硬编码的字符串数组，而是：
    1. knowledge_docs 表中的 active 文档；
    2. backend/uploads 中尚未登记到数据库的本地文件。

    数据库记录是在线增删改的主入口，file_path 只作为本地文件来源/存档。
    """

    SUPPORTED_FILE_SUFFIXES = {
        ".txt", ".md", ".markdown", ".csv", ".json", ".docx", ".xlsx"
    }
    TEXT_FILE_SUFFIXES = {".txt", ".md", ".markdown"}
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 150

    def __init__(self):
        self._initialized = False
        self.collection = None
        self.db_client = None
        self.model = None
        self.emb_fn = None
        self.documents: list[dict] = []
        self._index_lock = RLock()

        self.backend_dir = Path(__file__).resolve().parents[2]
        self.upload_dir = self._resolve_data_dir(settings.UPLOAD_DIR, "uploads")
        self.chroma_dir = self._resolve_data_dir(settings.CHROMA_DB_DIR, "chroma_db")
        self.model_path = settings.EMBEDDING_MODEL_PATH

        self.client = None
        self.device = None
        self.embed_model = settings.EMBEDDING_MODEL

    def _resolve_data_dir(self, configured_path: str, default_name: str) -> Path:
        """Resolve app data dirs relative to backend/ even when cwd is repo root."""
        raw = Path(configured_path or default_name)
        if raw.is_absolute():
            return raw

        backend_path = (self.backend_dir / raw).resolve()
        cwd_path = raw.resolve()
        if backend_path.exists() or not cwd_path.exists():
            return backend_path
        return cwd_path

    def _resolve_file_path(self, file_path: str | os.PathLike | None) -> Path | None:
        if not file_path:
            return None

        raw = Path(file_path)
        if raw.is_absolute():
            return raw

        candidates = [
            (self.backend_dir / raw).resolve(),
            (self.backend_dir.parent / raw).resolve(),
            raw.resolve(),
        ]
        for candidate in candidates:
            if candidate.exists():
                return candidate
        return candidates[0]

    def initialize(self, force_reload: bool = False):
        """同步初始化 Chroma，并从数据库/本地文件重建索引。"""
        start = time.perf_counter()
        with self._index_lock:
            if self._initialized and not force_reload:
                logger.info("rag initialize skipped already_initialized=true")
                return

            self.documents = self._load_source_documents()
            try:
                import chromadb
                from sentence_transformers import SentenceTransformer
                import torch

                if self.device is None:
                    self.device = "cuda" if torch.cuda.is_available() else "cpu"

                if self.emb_fn is None:
                    logger.info(
                        "rag loading embedding model path=%s device=%s docs=%s",
                        self.model_path,
                        self.device,
                        len(self.documents),
                    )
                    self.model = SentenceTransformer(
                        self.model_path,
                        device=self.device,
                        trust_remote_code=True,
                        model_kwargs={"attn_implementation": "sdpa"},
                    )
                    self.emb_fn = MyLocalQwenEF(self.model)

                os.makedirs(self.chroma_dir, exist_ok=True)
                self.db_client = chromadb.PersistentClient(path=str(self.chroma_dir))
                self.collection = self.db_client.get_or_create_collection(
                    name="LinXi_knowledge_base",
                    metadata={"hnsw:space": "cosine"},
                    embedding_function=self.emb_fn,
                )

                self._rebuild_index(self.documents)
                self._initialized = True
                logger.info(
                    "rag initialized docs=%s chunks=%s chroma_dir=%s duration_ms=%s",
                    len(self.documents),
                    self.collection.count(),
                    self.chroma_dir,
                    elapsed_ms(start),
                )
            except Exception as e:
                self.collection = None
                self._initialized = False
                logger.exception(
                    "rag initialize failed, fallback keyword only duration_ms=%s error=%s",
                    elapsed_ms(start),
                    e,
                )

    async def embed_batch(self, texts: list[str]) -> list[list[float]]:
        """批量向量化实现，保留给云端/HTTP embedding 方案使用。"""
        try:
            if self.client is None:
                from openai import AsyncOpenAI
                self.client = AsyncOpenAI(
                    api_key=settings.DASHSCOPE_API_KEY,
                    base_url=settings.VLLM_EMBED_BASE_URL,
                )
            response = await self.client.embeddings.create(
                model=self.embed_model,
                input=texts,
                dimensions=1024 if "v4" in self.embed_model.lower() else None,
                encoding_format="float",
            )
            return [item.embedding for item in response.data]
        except Exception as e:
            print(f"❌ 批量 Embedding 失败: {e}")
            raise e

    def _builtin_documents(self) -> list[dict]:
        return [
            {
                "id": doc.get("id", f"builtin_{index}"),
                "title": doc.get("title", ""),
                "category": doc.get("category", "general"),
                "content": doc.get("content", ""),
                "file_path": "",
                "source": "builtin",
                "db_id": None,
            }
            for index, doc in enumerate(LINGSHAN_KNOWLEDGE)
            if (doc.get("content") or "").strip()
        ]

    def _truncate_for_chat(self, doc: dict) -> dict:
        content = doc.get("content") or ""
        limit = settings.RAG_MAX_DOC_CHARS_FOR_CHAT
        if len(content) <= limit:
            return doc

        truncated = doc.copy()
        truncated["content"] = content[:limit]
        truncated["content_truncated"] = True
        truncated["original_content_chars"] = len(content)
        logger.warning(
            "rag document truncated for chat title=%s source=%s original_chars=%s kept_chars=%s file_path=%s",
            brief_text(doc.get("title"), 80),
            doc.get("source"),
            len(content),
            limit,
            doc.get("file_path", ""),
        )
        return truncated

    def _limit_documents_for_chat(self, docs: list[dict]) -> list[dict]:
        limited: list[dict] = []
        total = 0
        total_limit = settings.RAG_MAX_TOTAL_CHARS_FOR_CHAT
        for doc in docs:
            truncated = self._truncate_for_chat(doc)
            content_len = len(truncated.get("content") or "")
            if total + content_len > total_limit and limited:
                logger.warning(
                    "rag source documents total limit reached kept_docs=%s total_chars=%s total_limit=%s skipped_title=%s",
                    len(limited),
                    total,
                    total_limit,
                    brief_text(doc.get("title"), 80),
                )
                break
            limited.append(truncated)
            total += content_len
            if total >= total_limit:
                logger.warning(
                    "rag source documents total limit reached kept_docs=%s total_chars=%s total_limit=%s",
                    len(limited),
                    total,
                    total_limit,
                )
                break
        return limited

    def _load_source_documents(self) -> list[dict]:
        db_docs, known_file_paths = self._load_database_documents()
        file_docs = self._load_standalone_upload_documents(known_file_paths)
        builtin_docs = self._builtin_documents()
        docs = self._limit_documents_for_chat(db_docs + file_docs + builtin_docs)
        logger.info(
            "rag source documents loaded db_docs=%s file_docs=%s builtin_docs=%s total=%s",
            len(db_docs),
            len(file_docs),
            len(builtin_docs),
            len(docs),
        )
        return docs

    def _load_database_documents(self) -> tuple[list[dict], set[str]]:
        docs: list[dict] = []
        known_file_paths: set[str] = set()

        try:
            from app.models.database import KnowledgeDoc, SessionLocal
        except Exception as e:
            print(f"⚠️ 无法导入知识库数据库模型: {e}")
            return docs, known_file_paths

        db = SessionLocal()
        try:
            rows = db.query(KnowledgeDoc).all()
            for row in rows:
                file_path = self._resolve_file_path(row.file_path)
                if file_path:
                    known_file_paths.add(str(file_path.resolve()))

                if not row.is_active:
                    continue

                content = (row.content or "").strip()
                if not content and file_path and file_path.exists():
                    content = self.extract_file_content(
                        file_path,
                        max_chars=settings.RAG_MAX_DOC_CHARS_FOR_CHAT,
                    ).strip()
                if not content:
                    print(f"⚠️ 知识文档 {row.id} 内容为空，已跳过")
                    continue

                docs.append({
                    "id": self._source_id(row.id),
                    "title": (row.title or (file_path.stem if file_path else f"知识文档 {row.id}")).strip(),
                    "category": (row.category or "general").strip(),
                    "content": content,
                    "file_path": str(file_path) if file_path else "",
                    "source": "database",
                    "db_id": row.id,
                })
        except Exception as e:
            print(f"⚠️ 从数据库加载知识文档失败: {e}")
        finally:
            db.close()

        return docs, known_file_paths

    def _load_standalone_upload_documents(self, known_file_paths: set[str]) -> list[dict]:
        docs: list[dict] = []
        if not self.upload_dir.exists():
            return docs

        for file_path in sorted(self.upload_dir.rglob("*")):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in self.SUPPORTED_FILE_SUFFIXES:
                continue
            resolved = str(file_path.resolve())
            if resolved in known_file_paths:
                continue
            file_size = file_path.stat().st_size
            if (
                not settings.RAG_INITIALIZE_ON_CHAT
                and file_size > settings.RAG_MAX_UPLOAD_FILE_BYTES_FOR_CHAT
            ):
                logger.warning(
                    "rag upload file skipped for chat because too large path=%s size_bytes=%s limit_bytes=%s",
                    file_path,
                    file_size,
                    settings.RAG_MAX_UPLOAD_FILE_BYTES_FOR_CHAT,
                )
                continue

            try:
                content = self.extract_file_content(
                    file_path,
                    max_chars=settings.RAG_MAX_DOC_CHARS_FOR_CHAT,
                ).strip()
            except Exception as e:
                print(f"⚠️ 解析上传文件失败，已跳过 {file_path}: {e}")
                continue
            if not content:
                continue

            docs.append({
                "id": self._file_source_id(file_path),
                "title": file_path.stem,
                "category": "general",
                "content": content,
                "file_path": resolved,
                "source": "upload_file",
                "db_id": None,
            })
        return docs

    def extract_file_content(
        self,
        file_path: str | os.PathLike,
        max_chars: int | None = None,
    ) -> str:
        """Extract text from a supported knowledge file."""
        path = self._resolve_file_path(file_path)
        if path is None or not path.exists():
            raise FileNotFoundError(f"知识库文件不存在: {file_path}")

        suffix = path.suffix.lower()
        if suffix in self.TEXT_FILE_SUFFIXES:
            return self._limit_text(self._read_text_file(path), max_chars, path)
        if suffix == ".csv":
            return self._extract_csv(path, max_chars=max_chars)
        if suffix == ".json":
            return self._limit_text(self._extract_json(path), max_chars, path)
        if suffix == ".docx":
            return self._extract_docx(path, max_chars=max_chars)
        if suffix == ".xlsx":
            return self._extract_xlsx(path, max_chars=max_chars)
        raise ValueError(f"不支持的知识库文件格式: {suffix}")

    def _limit_text(self, text: str, max_chars: int | None, path: Path) -> str:
        if max_chars is None or len(text) <= max_chars:
            return text
        logger.warning(
            "rag file content limited path=%s original_chars=%s kept_chars=%s",
            path,
            len(text),
            max_chars,
        )
        return text[:max_chars]

    def _read_text_file(self, path: Path) -> str:
        last_error: Exception | None = None
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError as e:
                last_error = e
        if last_error:
            raise last_error
        return path.read_text(encoding="utf-8")

    def _append_limited_part(
        self,
        parts: list[str],
        value: str,
        max_chars: int | None,
        path: Path,
    ) -> bool:
        if not value:
            return False
        parts.append(value)
        if max_chars is not None and sum(len(part) + 1 for part in parts) >= max_chars:
            logger.warning("rag file parsing stopped at max_chars path=%s max_chars=%s", path, max_chars)
            return True
        return False

    def _join_limited_parts(self, parts: list[str], max_chars: int | None) -> str:
        text = "\n".join(parts)
        if max_chars is not None and len(text) > max_chars:
            return text[:max_chars]
        return text

    def _extract_csv(self, path: Path, max_chars: int | None = None) -> str:
        text = self._read_text_file(path)
        rows = []
        for row in csv.reader(text.splitlines()):
            values = [cell.strip() for cell in row if cell and cell.strip()]
            if values:
                if self._append_limited_part(rows, " | ".join(values), max_chars, path):
                    break
        return self._join_limited_parts(rows, max_chars)

    def _extract_json(self, path: Path) -> str:
        text = self._read_text_file(path)
        try:
            data = json.loads(text)
        except json.JSONDecodeError:
            return text
        return json.dumps(data, ensure_ascii=False, indent=2)

    def _extract_docx(self, path: Path, max_chars: int | None = None) -> str:
        try:
            from docx import Document
        except ImportError as e:
            raise RuntimeError("解析 .docx 需要安装 python-docx") from e

        document = Document(str(path))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                if self._append_limited_part(parts, text, max_chars, path):
                    return self._join_limited_parts(parts, max_chars)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    if self._append_limited_part(parts, " | ".join(cells), max_chars, path):
                        return self._join_limited_parts(parts, max_chars)
        return self._join_limited_parts(parts, max_chars)

    def _extract_xlsx(self, path: Path, max_chars: int | None = None) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as e:
            raise RuntimeError("解析 .xlsx 需要安装 openpyxl") from e

        workbook = load_workbook(str(path), read_only=True, data_only=True)
        parts: list[str] = []
        try:
            for sheet in workbook.worksheets:
                if self._append_limited_part(parts, f"【工作表：{sheet.title}】", max_chars, path):
                    break
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value).strip() for value in row if value not in (None, "")]
                    if values:
                        if self._append_limited_part(parts, " | ".join(values), max_chars, path):
                            return self._join_limited_parts(parts, max_chars)
        finally:
            workbook.close()
        return self._join_limited_parts(parts, max_chars)

    def _rebuild_index(self, docs: list[dict]):
        if self.collection is None:
            return
        self._clear_collection()
        self._add_documents_to_collection(docs)

    def _clear_collection(self):
        try:
            existing = self.collection.get()
            ids = existing.get("ids", []) if existing else []
            if ids:
                self.collection.delete(ids=ids)
        except Exception as e:
            print(f"⚠️ 清空旧 RAG 索引失败: {e}")

    def _add_documents_to_collection(self, docs: list[dict]):
        if self.collection is None:
            return

        ids: list[str] = []
        contents: list[str] = []
        metadatas: list[dict] = []

        for doc in docs:
            chunks = self._chunk_text(doc.get("content", ""))
            for index, chunk in enumerate(chunks):
                ids.append(f"{doc['id']}::chunk_{index:04d}")
                contents.append(chunk)
                metadatas.append({
                    "doc_id": doc["id"],
                    "title": doc.get("title", ""),
                    "category": doc.get("category", "general"),
                    "source": doc.get("source", "database"),
                    "file_path": doc.get("file_path", ""),
                    "chunk_index": index,
                })

        if not ids:
            return

        batch_size = 64
        for start in range(0, len(ids), batch_size):
            end = start + batch_size
            self.collection.add(
                ids=ids[start:end],
                documents=contents[start:end],
                metadatas=metadatas[start:end],
            )

    def _chunk_text(self, text: str) -> list[str]:
        text = re.sub(r"\n{3,}", "\n\n", (text or "").strip())
        if not text:
            return []
        if len(text) <= self.CHUNK_SIZE:
            return [text]

        chunks: list[str] = []
        current = ""
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        for paragraph in paragraphs:
            if len(paragraph) > self.CHUNK_SIZE:
                if current:
                    chunks.append(current.strip())
                    current = ""
                chunks.extend(self._split_long_text(paragraph))
                continue

            candidate = f"{current}\n\n{paragraph}" if current else paragraph
            if len(candidate) <= self.CHUNK_SIZE:
                current = candidate
            else:
                if current:
                    chunks.append(current.strip())
                current = paragraph

        if current:
            chunks.append(current.strip())
        return chunks

    def _split_long_text(self, text: str) -> list[str]:
        chunks: list[str] = []
        step = max(1, self.CHUNK_SIZE - self.CHUNK_OVERLAP)
        for start in range(0, len(text), step):
            chunk = text[start:start + self.CHUNK_SIZE].strip()
            if chunk:
                chunks.append(chunk)
        return chunks

    def _source_id(self, doc_id: str | int) -> str:
        value = str(doc_id)
        if value.startswith(("db_", "file_")):
            return value
        if value.isdigit():
            return f"db_{value}"
        return value

    def _file_source_id(self, file_path: Path) -> str:
        try:
            relative = file_path.relative_to(self.upload_dir)
        except ValueError:
            relative = file_path.name
        digest = hashlib.sha1(str(relative).encode("utf-8")).hexdigest()[:12]
        return f"file_{digest}"

    def safe_filename(self, filename: str) -> str:
        name = Path(filename or "knowledge").name
        stem = re.sub(r"[^\w.\-\u4e00-\u9fff]+", "_", name, flags=re.UNICODE).strip("._")
        return stem[:100] or "knowledge"

    def write_text_document_file(self, doc_id: str | int, title: str, content: str) -> str:
        os.makedirs(self.upload_dir, exist_ok=True)
        safe_title = self.safe_filename(title or f"knowledge_{doc_id}")
        file_path = self.upload_dir / f"db_{doc_id}_{safe_title}.md"
        file_path.write_text(content or "", encoding="utf-8")
        return str(file_path)

    def _delete_source(self, source_id: str):
        if self.collection is None:
            return
        try:
            existing = self.collection.get(where={"doc_id": source_id})
            ids = existing.get("ids", []) if existing else []
            if ids:
                self.collection.delete(ids=ids)
        except Exception as e:
            print(f"⚠️ 删除旧知识片段失败 {source_id}: {e}")

    def _upsert_snapshot(self, doc: dict):
        self.documents = [item for item in self.documents if item.get("id") != doc.get("id")]
        if (doc.get("content") or "").strip():
            self.documents.append(doc)

    def _build_runtime_doc(
        self,
        doc_id: str | int,
        title: str,
        category: str,
        content: str,
        file_path: str | None = None,
    ) -> dict:
        resolved_file = self._resolve_file_path(file_path)
        return {
            "id": self._source_id(doc_id),
            "title": title or "未命名知识文档",
            "category": category or "general",
            "content": content or "",
            "file_path": str(resolved_file) if resolved_file else "",
            "source": "database",
            "db_id": int(doc_id) if str(doc_id).isdigit() else None,
        }

    async def search(self, query: str, top_k: int = None) -> list[dict]:
        """语义检索；向量不可用时降级为基于数据库/文件快照的关键词检索。"""
        start = time.perf_counter()
        k = top_k or settings.RAG_TOP_K
        if not query or not query.strip():
            logger.info("rag search skipped empty query")
            return []

        if not self._initialized:
            if not settings.RAG_INITIALIZE_ON_CHAT:
                if not self.documents:
                    self.documents = self._load_source_documents()
                docs = self._keyword_search(query, k)
                logger.info(
                    "rag search keyword fallback because not initialized query=%s docs=%s top_score=%s duration_ms=%s",
                    brief_text(query, 120),
                    len(docs),
                    docs[0].get("score") if docs else None,
                    elapsed_ms(start),
                )
                return docs
            self.initialize()

        if self.collection is None:
            docs = self._keyword_search(query, k)
            logger.info(
                "rag search keyword fallback no collection query=%s docs=%s top_score=%s duration_ms=%s",
                brief_text(query, 120),
                len(docs),
                docs[0].get("score") if docs else None,
                elapsed_ms(start),
            )
            return docs

        try:
            with self._index_lock:
                count = self.collection.count()
                if count <= 0:
                    docs = self._keyword_search(query, k)
                    logger.info(
                        "rag search keyword fallback empty collection query=%s docs=%s top_score=%s duration_ms=%s",
                        brief_text(query, 120),
                        len(docs),
                        docs[0].get("score") if docs else None,
                        elapsed_ms(start),
                    )
                    return docs
                results = self.collection.query(
                    query_texts=[query],
                    n_results=min(k, count),
                )

            docs = []
            if results.get("documents") and results["documents"][0]:
                for i, content in enumerate(results["documents"][0]):
                    meta = results["metadatas"][0][i] if results.get("metadatas") else {}
                    dist = results["distances"][0][i] if results.get("distances") else 0.5
                    docs.append({
                        "id": meta.get("doc_id", ""),
                        "title": meta.get("title", ""),
                        "category": meta.get("category", ""),
                        "content": content,
                        "score": round(max(0.0, 1.0 - dist), 4),
                        "source": meta.get("source", ""),
                        "file_path": meta.get("file_path", ""),
                    })
            logger.info(
                "rag search vector done query=%s collection_count=%s docs=%s top_score=%s duration_ms=%s",
                brief_text(query, 120),
                count,
                len(docs),
                docs[0].get("score") if docs else None,
                elapsed_ms(start),
            )
            return docs
        except Exception as e:
            logger.exception(
                "rag search vector failed, fallback keyword query=%s duration_ms=%s error=%s",
                brief_text(query, 120),
                elapsed_ms(start),
                e,
            )
            docs = self._keyword_search(query, k)
            logger.info(
                "rag search keyword after vector failure docs=%s top_score=%s duration_ms=%s",
                len(docs),
                docs[0].get("score") if docs else None,
                elapsed_ms(start),
            )
            return docs

    async def search_local(self, query: str, top_k: int = None) -> list[dict]:
        """兼容旧调用。"""
        return await self.search(query, top_k)

    def _keyword_search(self, query: str, top_k: int) -> list[dict]:
        """关键词检索降级方案，使用当前数据库/文件快照。"""
        start = time.perf_counter()
        if not self.documents:
            self.documents = self._load_source_documents()

        scored = []
        query_lower = query.lower()
        terms = [term for term in re.split(r"\s+", query_lower) if term]

        category_keywords = {
            "history": ["历史", "渊源", "唐", "宋", "玄奘", "千年"],
            "attraction": ["景点", "大佛", "梵宫", "坛城", "塔", "寺"],
            "performance": ["表演", "演出", "九龙", "吉祥颂", "时间", "场次"],
            "route": ["路线", "游览", "推荐", "行程", "怎么玩"],
            "faq": ["门票", "价格", "开放", "时间", "停车", "注意"],
            "dining": ["餐饮", "吃饭", "素斋", "住宿", "酒店"],
            "nianhewan": ["拈花湾", "小镇", "花海", "灯光秀"],
            "temple": ["禅寺", "祥符", "撞钟", "银杏", "古井"],
            "general": [],
        }

        for doc in self.documents:
            content = (doc.get("content") or "")[:settings.RAG_MAX_DOC_CHARS_FOR_CHAT]
            title = doc.get("title", "")
            content_lower = content.lower()
            title_lower = title.lower()
            score = 0.0

            for term in terms:
                if term in title_lower:
                    score += 3.0
                if term in content_lower:
                    score += 1.5

            for char in query_lower:
                if char.strip() and char in title_lower:
                    score += 0.15
                if char.strip() and char in content_lower:
                    score += 0.05

            for cat, keywords in category_keywords.items():
                if doc.get("category") == cat and any(keyword in query for keyword in keywords):
                    score += 2.0

            if score > 0:
                scored.append((score, doc))

        scored.sort(key=lambda item: item[0], reverse=True)
        results = [
            {
                "id": doc.get("id", ""),
                "title": doc.get("title", ""),
                "category": doc.get("category", ""),
                "content": doc.get("content", ""),
                "score": round(min(1.0, score / 10), 4),
                "source": doc.get("source", ""),
                "file_path": doc.get("file_path", ""),
            }
            for score, doc in scored[:top_k]
        ]
        logger.debug(
            "rag keyword search done query=%s source_docs=%s matched=%s returned=%s duration_ms=%s",
            brief_text(query, 120),
            len(self.documents),
            len(scored),
            len(results),
            elapsed_ms(start),
        )
        return results

    def _index_knowledge(self, docs: list[dict]):
        """兼容旧方法：将传入文档直接写入当前向量库。"""
        normalized = []
        for doc in docs:
            normalized.append({
                "id": str(doc.get("id")),
                "title": doc.get("title", ""),
                "category": doc.get("category", "general"),
                "content": doc.get("content", ""),
                "file_path": doc.get("file_path", ""),
                "source": doc.get("source", "manual"),
                "db_id": doc.get("db_id"),
            })
        self._add_documents_to_collection(normalized)

    def _index_knowledge_local(self, docs: list[dict]):
        """兼容旧方法。"""
        self._index_knowledge(docs)

    async def add_document(
        self,
        doc_id: str | int,
        title: str,
        category: str,
        content: str,
        file_path: str | None = None,
    ):
        """在线新增文档，并同步写入 Chroma。"""
        doc = self._build_runtime_doc(doc_id, title, category, content, file_path)
        with self._index_lock:
            if not self._initialized:
                self.initialize()
            self._upsert_snapshot(doc)
            if self.collection is None:
                return
            self._delete_source(doc["id"])
            self._add_documents_to_collection([doc])

    async def update_document(
        self,
        doc_id: str | int,
        title: str,
        category: str,
        content: str,
        file_path: str | None = None,
    ):
        """在线更新文档，等价于删除旧片段后重新索引。"""
        await self.add_document(doc_id, title, category, content, file_path)

    async def delete_document(self, doc_id: str | int):
        """在线删除文档的所有向量片段。"""
        source_id = self._source_id(doc_id)
        with self._index_lock:
            if not self._initialized:
                self.initialize()
            self.documents = [doc for doc in self.documents if doc.get("id") != source_id]
            self._delete_source(source_id)

    async def reload_from_database(self):
        """重新从数据库和 backend/uploads 文件扫描并重建索引。"""
        with self._index_lock:
            if not self._initialized:
                self.initialize(force_reload=True)
                return
            self.documents = self._load_source_documents()
            self._rebuild_index(self.documents)

    def get_all_documents(self) -> list[dict]:
        """获取当前数据库/文件知识快照。"""
        if not self.documents:
            self.documents = self._load_source_documents()
        return [doc.copy() for doc in self.documents]

    def format_context(self, docs: list[dict], max_chars: int = 2000) -> str:
        """将检索结果格式化为 Prompt 上下文。"""
        if not docs:
            return "（未检索到相关景区知识，请根据通用佛教文化知识回答）"
        parts = []
        total = 0
        for doc in docs:
            title = doc.get("title", "未命名知识")
            score = doc.get("score", 0)
            content = doc.get("content", "")
            snippet = f"【{title}】（相关度{score:.0%}）\n{content}"
            if total + len(snippet) > max_chars:
                break
            parts.append(snippet)
            total += len(snippet)
        return "\n\n---\n\n".join(parts)

# 全局单例
rag_service = RAGService()
